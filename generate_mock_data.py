import os
import json
import docx
from pathlib import Path
import random

def generate():
    data_dir = Path("data")
    data_dir.mkdir(exist_ok=True)
    
    # 1. Generate job_description.docx
    doc = docx.Document()
    doc.add_heading("Job Description: Senior Machine Learning Engineer", level=1)
    doc.add_paragraph("Role: Senior Machine Learning Engineer")
    doc.add_paragraph("Location: Bengaluru, India / Remote")
    doc.add_paragraph("Experience: 5-9 years")
    doc.add_paragraph("Required Skills: Python, LLMs, RAG, Embeddings, FAISS, FastAPI, PyTorch")
    doc.add_paragraph("Preferred Skills: LangChain, MLOps, Docker, AWS")
    doc.add_paragraph("We are looking for an expert in Search and AI Retrieval Engineering who can design and build our next-generation candidate matching systems using vector search, neural reranking, and semantic indexes.")
    doc.save(data_dir / "job_description.docx")
    print("Generated data/job_description.docx")

    # 2. Generate candidates.jsonl programmatically (50 mock candidates)
    skills_pool = [
        "Python", "LLMs", "RAG", "Embeddings", "FAISS", "FastAPI", "PyTorch",
        "LangChain", "MLOps", "Docker", "AWS", "SQL", "Spark", "TensorFlow", "Scikit-Learn"
    ]
    
    headlines = [
        "Senior Machine Learning Engineer", "ML Engineer", "Data Scientist", 
        "Software Engineer (AI)", "NLP Research Engineer", "AI Platform Engineer"
    ]
    
    schools = [
        ("IIT Bombay", "tier_1"), ("IIT Delhi", "tier_1"), ("IISc Bangalore", "tier_1"),
        ("BITS Pilani", "tier_2"), ("IIIT Hyderabad", "tier_2"),
        ("Local Engineering College", "tier_3"), ("Unknown University", "unknown")
    ]

    candidates = []
    for i in range(1, 51):
        candidate_id = f"CAND_{100000 + i}"
        headline = random.choice(headlines)
        
        # Decide skills list
        num_skills = random.randint(3, 8)
        selected_skills = random.sample(skills_pool, num_skills)
        candidate_skills = []
        for s in selected_skills:
            candidate_skills.append({
                "name": s,
                "duration_months": random.randint(3, 60),
                "endorsements": random.randint(0, 15)
            })
            
        # Career History
        num_jobs = random.randint(1, 3)
        career_history = []
        for j in range(num_jobs):
            career_history.append({
                "title": random.choice(headlines),
                "description": f"Worked on machine learning models, search indices, and API servers using {', '.join(selected_skills[:2])}."
            })
            
        school, tier = random.choice(schools)
        education = [{"tier": tier, "degree": "B.Tech / M.Tech in CS", "school": school}]
        
        redrob_signals = {
            "recruiter_response_rate": round(random.uniform(0.3, 1.0), 2),
            "interview_completion_rate": round(random.uniform(0.3, 1.0), 2),
            "offer_acceptance_rate": round(random.uniform(0.3, 1.0), 2),
            "open_to_work_flag": random.choice([True, False]),
            "applications_submitted_30d": random.randint(0, 20),
            "willing_to_relocate": random.choice([True, False]),
            "github_activity_score": random.randint(10, 100),
            "profile_completeness_score": random.randint(40, 100),
            "search_appearance_30d": random.randint(5, 400)
        }
        
        candidates.append({
            "candidate_id": candidate_id,
            "profile": {
                "headline": headline,
                "summary": f"Experienced professional specialized in AI and backend architectures using {', '.join(selected_skills)}."
            },
            "skills": candidate_skills,
            "career_history": career_history,
            "education": education,
            "redrob_signals": redrob_signals
        })

    with open(data_dir / "candidates.jsonl", "w") as f:
        for c in candidates:
            f.write(json.dumps(c) + "\n")
    print("Generated data/candidates.jsonl")

if __name__ == "__main__":
    generate()
